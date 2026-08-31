from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins to 1 inch
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Name
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run("FIRSTNAME LASTNAME")
name_run.bold = True
name_run.font.size = Pt(24)
name_run.font.name = 'Arial'

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run("City, State - (555) 555-5555 - email@example.com - linkedin.com/in/username")
contact_run.font.size = Pt(11)
contact_run.font.name = 'Arial'

doc.add_paragraph() # Spacing

# Professional Summary
summary_heading = doc.add_heading("PROFESSIONAL SUMMARY", level=1)
for run in summary_heading.runs:
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.bold = True

summary_body = doc.add_paragraph("Results-driven professional with X years of experience in [Industry]. Proven track record of [Key Achievement 1] and [Key Achievement 2]. Skilled in [Skill 1], [Skill 2], and [Skill 3].")
summary_body.runs[0].font.name = 'Arial'
summary_body.runs[0].font.size = Pt(11)

# Work Experience
exp_heading = doc.add_heading("PROFESSIONAL EXPERIENCE", level=1)
for run in exp_heading.runs:
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.bold = True

# Job 1
job1 = doc.add_paragraph()
j1_title = job1.add_run("Job Title")
j1_title.bold = True
j1_title.font.name = 'Arial'
j1_title.font.size = Pt(11)
job1.add_run(" | ").font.name = 'Arial'
j1_company = job1.add_run("Company Name, City, State")
j1_company.font.name = 'Arial'
j1_company.font.size = Pt(11)
job1.add_run("\t\t\t\t\t\t\t\t\t\tMonth Year - Present").font.name = 'Arial' # Hacky right align for simplicity

j1_bullet1 = doc.add_paragraph("Spearheaded [Project/Initiative], resulting in a [Number]% increase in [Metric] over [Timeframe].", style='List Bullet')
j1_bullet1.runs[0].font.name = 'Arial'
j1_bullet1.runs[0].font.size = Pt(11)
j1_bullet2 = doc.add_paragraph("Managed a cross-functional team of [Number] to deliver [Product/Service], cutting costs by $[Amount].", style='List Bullet')
j1_bullet2.runs[0].font.name = 'Arial'
j1_bullet2.runs[0].font.size = Pt(11)
j1_bullet3 = doc.add_paragraph("Collaborated with [Department] to streamline [Process], reducing manual processing time by [Number] hours per week.", style='List Bullet')
j1_bullet3.runs[0].font.name = 'Arial'
j1_bullet3.runs[0].font.size = Pt(11)

doc.add_paragraph()

# Education
edu_heading = doc.add_heading("EDUCATION", level=1)
for run in edu_heading.runs:
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.bold = True

edu = doc.add_paragraph()
e_degree = edu.add_run("Degree Name (e.g., Bachelor of Science in Computer Science)")
e_degree.bold = True
e_degree.font.name = 'Arial'
e_degree.font.size = Pt(11)
edu.add_run("\n").font.name = 'Arial'
e_school = edu.add_run("University Name, City, State")
e_school.font.name = 'Arial'
e_school.font.size = Pt(11)

doc.add_paragraph()

# Skills
skills_heading = doc.add_heading("SKILLS", level=1)
for run in skills_heading.runs:
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.bold = True

skills = doc.add_paragraph()
s_run = skills.add_run("Technical Skills: Skill 1, Skill 2, Skill 3, Skill 4\nSoft Skills: Skill A, Skill B, Skill C")
s_run.font.name = 'Arial'
s_run.font.size = Pt(11)

doc.save("GetATSReady_Template.docx")
print("Template created!")
